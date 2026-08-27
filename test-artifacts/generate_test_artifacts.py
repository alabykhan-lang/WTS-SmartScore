"""Generate invented SmartScore paper and export fixtures.

The QR is intentionally in the top-centre header zone. It carries page identity
only; no handwritten score values are encoded.
"""
from pathlib import Path
import json
import sys
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.graphics.barcode.qr import QrCodeWidget
from reportlab.graphics.shapes import Drawing
from reportlab.graphics import renderPDF
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PDF_DIR = ROOT / "output" / "pdf"
DOCX_DIR = ROOT / "output" / "docx"
OCR_DIR = ROOT / "output" / "ocr"
for directory in (PDF_DIR, DOCX_DIR, OCR_DIR):
    directory.mkdir(parents=True, exist_ok=True)

NAVY = colors.HexColor("#0B2545")
BLUE = colors.HexColor("#1E5AA8")
PALE = colors.HexColor("#EAF1F8")
GRID = colors.HexColor("#AEBCCD")


def draw_qr(c, payload, center_x, bottom_y, size=24 * mm):
    c.setFillColor(colors.white)
    c.rect(center_x - size / 2 - 3 * mm, bottom_y - 3 * mm, size + 6 * mm, size + 6 * mm, fill=1, stroke=0)
    widget = QrCodeWidget(payload)
    x1, y1, x2, y2 = widget.getBounds()
    drawing = Drawing(size, size, transform=[size / (x2 - x1), 0, 0, size / (y2 - y1), -x1 * size / (x2 - x1), -y1 * size / (y2 - y1)])
    drawing.add(widget)
    renderPDF.draw(drawing, c, center_x - size / 2, bottom_y)


def header(c, title, class_name, subject, sheet_id, page_id, page_number, layout_id):
    width, height = landscape(A4)
    c.setFillColor(NAVY)
    c.rect(0, height - 36 * mm, width, 36 * mm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(18 * mm, height - 12 * mm, "WAY TO SUCCESS STANDARD SCHOOLS")
    c.setFont("Helvetica-Bold", 14)
    c.drawString(18 * mm, height - 21 * mm, title)
    c.setFont("Helvetica", 8.5)
    c.drawString(18 * mm, height - 29 * mm, f"Class: {class_name}    Subject group: {subject}    Session: 2026/2027    Term: FIRST")
    payload = json.dumps({"v": 1, "s": sheet_id, "p": page_id, "n": page_number, "l": layout_id, "c": class_name, "u": subject, "t": "FIRST"}, separators=(",", ":"))
    draw_qr(c, payload, width / 2, height - 24 * mm, 20 * mm)
    c.setFillColor(colors.HexColor("#D9E7F5"))
    c.setFont("Helvetica", 6.5)
    c.drawCentredString(width / 2, height - 34 * mm, f"{page_id}  •  scan identity only")


def secondary_page(c, sheet_id, page_id, page_number, class_name, subject, names):
    width, height = landscape(A4)
    header(c, "SMARTSCORE BROADSHEET", class_name, subject, sheet_id, page_id, page_number, "secondary-single-subject-v3")
    x0 = 15 * mm
    y_top = height - 47 * mm
    col_widths = [13 * mm, 92 * mm, 27 * mm, 27 * mm, 27 * mm, 33 * mm]
    labels = ["No.", "Student Name", "CA1 / 10", "CA2 / 10", "CA3 / 10", "EXAM / 70"]
    c.setFillColor(PALE)
    c.rect(x0, y_top - 10 * mm, sum(col_widths), 10 * mm, fill=1, stroke=0)
    c.setFillColor(NAVY)
    c.setFont("Helvetica-Bold", 7.5)
    x = x0
    for label, col_width in zip(labels, col_widths):
        c.drawCentredString(x + col_width / 2, y_top - 6.2 * mm, label) if label != "Student Name" else c.drawString(x + 2 * mm, y_top - 6.2 * mm, label)
        x += col_width
    row_h = 10 * mm
    c.setFont("Helvetica", 7.5)
    for row, name in enumerate(names, 1):
        y = y_top - 10 * mm - row * row_h
        if row % 2 == 0:
            c.setFillColor(colors.HexColor("#F8FAFC"))
            c.rect(x0, y, sum(col_widths), row_h, fill=1, stroke=0)
        c.setFillColor(NAVY)
        x = x0
        values = [str(row + (page_number - 1) * 15), name, "", "", "", ""]
        for index, (value, col_width) in enumerate(zip(values, col_widths)):
            c.drawCentredString(x + col_width / 2, y + 3.4 * mm, value) if index != 1 else c.drawString(x + 2 * mm, y + 3.4 * mm, value)
            x += col_width
    c.setStrokeColor(GRID)
    x = x0
    for col_width in col_widths:
        c.line(x, y_top, x, y_top - 10 * mm - len(names) * row_h)
        x += col_width
    c.line(x, y_top, x, y_top - 10 * mm - len(names) * row_h)
    for row in range(len(names) + 2):
        y = y_top - row * row_h
        c.line(x0, y, x0 + sum(col_widths), y)
    c.setFillColor(colors.HexColor("#536579"))
    c.setFont("Helvetica-Oblique", 6.5)
    c.drawString(x0, 11 * mm, "Write scores inside the printed cells. QR remains in the header and outside score ROIs.")
    c.showPage()


def generate_secondary(path, sheet_id, class_name, subject, population, page_count):
    c = canvas.Canvas(str(path), pagesize=landscape(A4))
    names = [
        "Adigun Bazim", "Bakare Fathiat", "Oyelami Muiz", "Adam Kashfat", "Usman Toheebat",
        "Hassan Ibrahim", "Adebayo Mariam", "Ojo Samuel", "Akinola Temilade", "Lawal Hammed",
        "Oyediran Zainab", "Afolabi Daniel", "Salami Khadijat", "Olatunji Victor", "Adeleke Rofiat",
        "Ibrahim Sodiq", "Ajibade Deborah", "Amoo Ridwan", "Oyeniyi Faruq", "Balogun Aminat",
        "Oluwaseun Martins", "Adeniyi Barakat", "Ogundele Praise", "Yusuf Lateefah", "Fashina Habeeb",
        "Oladapo Faith", "Aderibigbe Halimat", "Ogunleye Malik", "Taiwo Kemi", "Kareem Mubashir",
        "Adekunle Seyi", "Bello Mariam", "Ogunbiyi Peter", "Sanni Rukayat", "Adewale Tobi", "Ishola Zainab"
    ][:population]
    for page_number in range(1, page_count + 1):
        start = (page_number - 1) * 15
        secondary_page(c, sheet_id, f"{sheet_id}-P{page_number}", page_number, class_name, subject, names[start:start + 15])
    c.save()


def generate_primary(path):
    c = canvas.Canvas(str(path), pagesize=landscape(A4))
    width, height = landscape(A4)
    sheet_id = "WTS-SS-PRIMARY-MULTI-001"
    page_id = f"{sheet_id}-P1"
    header(c, "SMARTSCORE MULTI-SUBJECT BROADSHEET", "TEST PRIMARY 4", "English • Mathematics • Basic Science • Social Studies", sheet_id, page_id, 1, "primary-four-subject-v1")
    x0 = 10 * mm
    y_top = height - 47 * mm
    columns = [("No.", 11 * mm), ("Student Name", 55 * mm)]
    for subject in ["ENGLISH", "MATHEMATICS", "BASIC SCIENCE", "SOCIAL STUDIES"]:
        columns.extend([(subject, 17 * mm), ("", 17 * mm)])
    group_width = 34 * mm
    c.setFillColor(PALE)
    c.rect(x0, y_top - 17 * mm, sum(width for _, width in columns), 17 * mm, fill=1, stroke=0)
    c.setFillColor(NAVY)
    c.setFont("Helvetica-Bold", 5.8)
    x = x0
    c.drawCentredString(x + columns[0][1] / 2, y_top - 10 * mm, "No.")
    x += columns[0][1]
    c.drawString(x + 2 * mm, y_top - 10 * mm, "Student Name")
    x += columns[1][1]
    subjects = ["ENGLISH", "MATHEMATICS", "BASIC SCIENCE", "SOCIAL STUDIES"]
    for subject_index, subject in enumerate(subjects):
        span = group_width
        c.drawCentredString(x + span / 2, y_top - 6 * mm, subject)
        c.drawCentredString(x + 17 * mm / 2, y_top - 13 * mm, "CA")
        c.drawCentredString(x + 17 * mm + 17 * mm / 2, y_top - 13 * mm, "EXAM")
        x += span
    names = ["Aisha Bello", "David Adeyemi", "Favour Okoro", "Mariam Yusuf", "Peter Akin", "Tolu Ajayi", "Zainab Sanni", "Samuel Ojo", "Esther Lawal", "Daniel Musa", "Hauwa Ibrahim", "Michael Adewale"]
    row_h = 10 * mm
    c.setFont("Helvetica", 6.2)
    original_columns = [("No.", 11 * mm), ("Student Name", 55 * mm)] + [("", 17 * mm)] * 8
    for row, name in enumerate(names, 1):
        y = y_top - 17 * mm - row * row_h
        if row % 2 == 0:
            c.setFillColor(colors.HexColor("#F8FAFC"))
            c.rect(x0, y, sum(w for _, w in original_columns), row_h, fill=1, stroke=0)
        c.setFillColor(NAVY)
        x = x0
        for idx, (_, col_width) in enumerate(original_columns):
            value = str(row) if idx == 0 else name if idx == 1 else ""
            c.drawCentredString(x + col_width / 2, y + 3.4 * mm, value) if idx != 1 else c.drawString(x + 2 * mm, y + 3.4 * mm, value)
            x += col_width
    c.setStrokeColor(GRID)
    x = x0
    for _, col_width in original_columns:
        c.line(x, y_top, x, y_top - 17 * mm - len(names) * row_h)
        x += col_width
    c.line(x, y_top, x, y_top - 17 * mm - len(names) * row_h)
    for row in range(len(names) + 3): c.line(x0, y_top - row * row_h, x0 + sum(w for _, w in original_columns), y_top - row * row_h)
    c.setFillColor(colors.HexColor("#536579"))
    c.setFont("Helvetica-Oblique", 6.5)
    c.drawString(x0, 11 * mm, "Primary layout test: four subject groups share one page; each score ROI stays independent.")
    c.showPage()
    c.save()


def generate_script_pdf(path):
    c = canvas.Canvas(str(path), pagesize=A4)
    for number, text in [(1, "STUDENT NAME: ADEBAYO KHALID\nADMISSION NO: TEST-SS1-001\nCLASS: TEST SS1\nSUBJECT: ECONOMICS\nTERMINAL EXAMINATION"), (2, "Question 4. Explain the law of demand.\n\nOCR is a derived convenience; the corrected page image remains authoritative.")]:
        width, height = A4
        c.setFillColor(NAVY); c.rect(0, height - 30 * mm, width, 30 * mm, fill=1, stroke=0)
        c.setFillColor(colors.white); c.setFont("Helvetica-Bold", 16); c.drawString(18 * mm, height - 16 * mm, f"SMARTSCORE SCRIPT EXPORT • PAGE {number}")
        c.setFillColor(NAVY); c.setFont("Helvetica", 11)
        y = height - 55 * mm
        for line in text.split("\n"):
            c.drawString(18 * mm, y, line); y -= 9 * mm
        c.showPage()
    c.save()


def generate_docx(path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1); section.right_margin = Inches(1); section.bottom_margin = Inches(1); section.left_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor(0, 0, 0)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("SmartScore Script Export")
    run.bold = True; run.font.size = Pt(23); run.font.color.rgb = RGBColor(11, 37, 69)
    doc.add_paragraph("ADEBAYO KHALID  •  ECONOMICS  •  TEST SS1")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    for row, (label, value) in zip(table.rows, [("Script ID", "TEST-SCRIPT-001"), ("Pages", "2"), ("Identity", "OCR suggestion; reviewable"), ("Source", "Corrected page images")]):
        row.cells[0].text = label; row.cells[1].text = value
    doc.add_heading("Page 1 OCR", level=1)
    doc.add_paragraph("STUDENT NAME: ADEBAYO KHALID\nADMISSION NO: TEST-SS1-001\nCLASS: TEST SS1\nSUBJECT: ECONOMICS")
    doc.add_page_break()
    doc.add_heading("Page 2 OCR", level=1)
    doc.add_paragraph("Question 4. Explain the law of demand.")
    doc.add_paragraph("This Word export is a best-effort OCR/layout reconstruction. Original corrected images remain authoritative.")
    doc.save(str(path))


def generate_pdfs():
    generate_secondary(PDF_DIR / "secondary-single-subject-8-students.pdf", "WTS-SS-SECONDARY-ONE-001", "TEST JSS1", "English Language", 8, 1)
    generate_secondary(PDF_DIR / "secondary-large-36-students-3-pages.pdf", "WTS-SS-SECONDARY-LARGE-001", "TEST SS2", "Mathematics", 36, 3)
    generate_primary(PDF_DIR / "primary-four-subject-one-page.pdf")
    generate_script_pdf(PDF_DIR / "script-export-example.pdf")


def generate_docx_and_ocr():
    generate_docx(DOCX_DIR / "script-export-example.docx")
    (OCR_DIR / "script-export-example.txt").write_text("===== PAGE 1 =====\nSTUDENT NAME: ADEBAYO KHALID\nADMISSION NO: TEST-SS1-001\nCLASS: TEST SS1\nSUBJECT: ECONOMICS\n\n===== PAGE 2 =====\nQuestion 4. Explain the law of demand.\n", encoding="utf-8")
    (OCR_DIR / "script-export-example.json").write_text(json.dumps({"script_id": "TEST-SCRIPT-001", "pages": [{"page_number": 1, "image": "pages/page-001.jpg", "text": "STUDENT NAME: ADEBAYO KHALID"}, {"page_number": 2, "image": "pages/page-002.jpg", "text": "Question 4. Explain the law of demand."}]}, indent=2), encoding="utf-8")


def main():
    option = sys.argv[1] if len(sys.argv) > 1 else "all"
    if option in ("all", "pdf"):
        generate_pdfs()
    if option in ("all", "docx"):
        generate_docx_and_ocr()


if __name__ == "__main__":
    main()
